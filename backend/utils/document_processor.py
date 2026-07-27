"""Document processing for text extraction and OpenAI visual descriptions."""
from __future__ import annotations

import base64
import io
import logging
import re
import tempfile
from dataclasses import dataclass
from pathlib import Path
from typing import Any

import PyPDF2
from docx import Document
from openai import OpenAI
from PIL import Image

Image.MAX_IMAGE_PIXELS = 40_000_000
logger = logging.getLogger(__name__)

from ..config import (
    MAX_VISUAL_IMAGES,
    OPENAI_IMAGE_DETAIL,
    OPENAI_MODEL,
    SUPPORTED_EXTENSIONS,
    VISUAL_DESCRIPTION_MAX_TOKENS,
    require_openai_api_key,
)


class DocumentProcessingError(Exception):
    """Raised when a file cannot produce usable evidence."""


@dataclass
class _PdfEmbeddedImage:
    page: int
    image_index: int
    image: Image.Image
    name: str | None = None


class DocumentProcessor:
    """Convert uploaded files into retrieval-ready text and visual content items."""

    def __init__(self, client: OpenAI | None = None):
        self.client = client

    def process_file(self, file_path: Path, document_id: str, filename: str) -> dict[str, Any]:
        ext = file_path.suffix.lower()
        if ext not in SUPPORTED_EXTENSIONS:
            raise DocumentProcessingError(f"Unsupported file format: {ext}")

        if ext == ".pdf":
            return self._process_pdf(file_path, document_id, filename)
        if ext == ".docx":
            return self._process_docx(file_path, document_id, filename)
        return self._process_image(file_path, document_id, filename)

    def _base_metadata(
        self,
        document_id: str,
        filename: str,
        file_type: str,
        content_type: str,
        page: int | None = None,
    ) -> dict[str, Any]:
        return {
            "document_id": document_id,
            "filename": filename,
            "file_type": file_type,
            "page": page,
            "content_type": content_type,
        }

    def _process_pdf(self, file_path: Path, document_id: str, filename: str) -> dict[str, Any]:
        items: list[dict[str, Any]] = []
        warnings: list[str] = []
        embedded_images: list[_PdfEmbeddedImage] = []
        pages = 0
        pdf_image_scan_enabled = False

        with file_path.open("rb") as handle:
            reader = PyPDF2.PdfReader(handle)
            pages = len(reader.pages)
            for index, page in enumerate(reader.pages, start=1):
                text = (page.extract_text() or "").strip()
                if text:
                    items.append(
                        {
                            "text": f"PDF text evidence from {filename}, page {index}:\n{text}",
                            "metadata": self._base_metadata(document_id, filename, "pdf", "text", index),
                        }
                    )
                    inferred_tables = self._infer_tables_from_page_text(text)
                    for table_index, inferred_table in enumerate(inferred_tables, start=1):
                        metadata = self._base_metadata(document_id, filename, "pdf", "table", index)
                        metadata["table_index"] = f"text-{table_index}"
                        items.append(
                            {
                                "text": (
                                    f"PDF table-like evidence inferred from text in {filename}, page {index}. "
                                    "Use this evidence for questions about tables, rows, columns, entries, labels, "
                                    "metrics, percentages, numeric values, and captions.\n"
                                    f"{inferred_table}"
                                ),
                                "metadata": metadata,
                            }
                        )

            if pages:
                pdf_image_scan_enabled = MAX_VISUAL_IMAGES > 0
                embedded_images, image_warnings = self._extract_pdf_embedded_images(reader, MAX_VISUAL_IMAGES)
                warnings.extend(image_warnings)

        table_items, table_warnings = self._extract_pdf_tables(file_path, document_id, filename)
        items.extend(table_items)
        warnings.extend(table_warnings)

        if embedded_images:
            for embedded_image in embedded_images:
                try:
                    description = self._describe_image(
                        embedded_image.image,
                        context=(
                            f"This is embedded image {embedded_image.image_index} extracted directly from page "
                            f"{embedded_image.page} of PDF file {filename}. Extract visible evidence for retrieval. "
                            "If there is a figure, chart, pie chart, table, leaderboard, caption, or title, include "
                            "its exact visible label such as Figure 2 or Table 1 when present."
                        ),
                    )
                except Exception as exc:
                    logger.warning(
                        "Visual analysis failed for document_id=%s page=%s embedded_image=%s",
                        document_id,
                        embedded_image.page,
                        embedded_image.image_index,
                        exc_info=exc,
                    )
                    if self._is_openai_configuration_or_quota_error(exc):
                        warnings.append(self._public_openai_visual_warning(exc))
                        break
                    warnings.append(
                        f"Visual analysis failed for embedded image {embedded_image.image_index} "
                        f"on page {embedded_image.page}."
                    )
                    continue
                if description:
                    metadata = self._base_metadata(document_id, filename, "pdf", "visual", embedded_image.page)
                    metadata["image_index"] = embedded_image.image_index
                    if embedded_image.name:
                        metadata["image_name"] = embedded_image.name
                    items.append(
                        {
                            "text": (
                                f"PDF embedded image visual evidence from {filename}, page {embedded_image.page}, "
                                f"image {embedded_image.image_index}:\n{description}"
                            ),
                            "metadata": metadata,
                        }
                    )
        elif pages and pdf_image_scan_enabled:
            warnings.append("No embedded images were found in this PDF.")

        if not items:
            raise DocumentProcessingError("No extractable text or visual evidence was found in this PDF.")

        return self._result(document_id, filename, "pdf", pages, items, warnings)

    def _extract_pdf_embedded_images(
        self,
        reader: Any,
        image_limit: int,
    ) -> tuple[list[_PdfEmbeddedImage], list[str]]:
        embedded_images: list[_PdfEmbeddedImage] = []
        warnings: list[str] = []
        skipped_images = 0

        if image_limit <= 0:
            warnings.append("PDF visual analysis is disabled because MAX_VISUAL_IMAGES is 0.")
            return embedded_images, warnings

        for page_index, page in enumerate(reader.pages):
            page_number = page_index + 1

            try:
                page_images = list(page.images or [])
            except AttributeError:
                warnings.append("PDF embedded image extraction requires PyPDF2 3.0.0 or newer.")
                break
            except Exception as exc:
                logger.warning("PDF embedded image extraction failed for page=%s", page_number, exc_info=exc)
                warnings.append(f"Embedded image extraction failed for page {page_number}.")
                continue

            for image_index, image_file in enumerate(page_images, start=1):
                try:
                    image = self._open_pdf_embedded_image(image_file)
                except Exception as exc:
                    logger.warning(
                        "PDF embedded image decoding failed for page=%s image=%s",
                        page_number,
                        image_index,
                        exc_info=exc,
                    )
                    skipped_images += 1
                    continue
                if image is None:
                    logger.warning(
                        "Unable to decode embedded image on page=%s image=%s",
                        page_number,
                        image_index,
                    )
                    skipped_images += 1
                    continue
                embedded_images.append(
                    _PdfEmbeddedImage(
                        page=page_number,
                        image_index=image_index,
                        image=image,
                        name=self._pdf_embedded_image_name(image_file),
                    )
                )
                if len(embedded_images) >= image_limit:
                    warnings.append(
                        f"Visual analysis was limited to the first {image_limit} embedded images found in this PDF."
                    )
                    if skipped_images:
                        warnings.append(f"{skipped_images} embedded PDF image(s) could not be decoded.")
                    return embedded_images, warnings

        if skipped_images:
            warnings.append(f"{skipped_images} embedded PDF image(s) could not be decoded.")
        return embedded_images, warnings

    def _open_pdf_embedded_image(self, image_file: Any) -> Image.Image | None:
        image = getattr(image_file, "image", None)
        if isinstance(image, Image.Image):
            image.load()
            return image.copy()

        data = getattr(image_file, "data", None)
        if not data:
            return None

        with Image.open(io.BytesIO(data)) as opened_image:
            opened_image.load()
            return opened_image.copy()

    def _pdf_embedded_image_name(self, image_file: Any) -> str | None:
        name = getattr(image_file, "name", None)
        return str(name) if name else None

    def _process_docx(self, file_path: Path, document_id: str, filename: str) -> dict[str, Any]:
        doc = Document(str(file_path))
        items: list[dict[str, Any]] = []
        warnings: list[str] = []

        paragraphs = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
        if paragraphs:
            items.append(
                {
                    "text": "\n".join(paragraphs),
                    "metadata": self._base_metadata(document_id, filename, "docx", "text"),
                }
            )

        image_count = 0
        for rel in doc.part.rels.values():
            if "image" not in rel.reltype:
                continue
            if MAX_VISUAL_IMAGES <= 0:
                warnings.append("DOCX visual analysis is disabled because MAX_VISUAL_IMAGES is 0.")
                break
            if image_count >= MAX_VISUAL_IMAGES:
                warnings.append(
                    f"Visual analysis was limited to the first {MAX_VISUAL_IMAGES} embedded images found in this DOCX."
                )
                break
            image_count += 1
            try:
                with Image.open(io.BytesIO(rel.target_part.blob)) as opened_image:
                    opened_image.load()
                    image = opened_image.copy()
                description = self._describe_image(
                    image,
                    context=f"This is embedded image {image_count} from DOCX file {filename}.",
                )
            except Exception as exc:
                logger.warning(
                    "Visual analysis failed for document_id=%s embedded_image=%s",
                    document_id,
                    image_count,
                    exc_info=exc,
                )
                if self._is_openai_configuration_or_quota_error(exc):
                    warnings.append(self._public_openai_visual_warning(exc))
                    break
                warnings.append(f"Visual analysis failed for embedded image {image_count}.")
                continue
            if description:
                metadata = self._base_metadata(document_id, filename, "docx", "visual")
                metadata["image_index"] = image_count
                items.append({"text": f"DOCX visual evidence from {filename}, image {image_count}:\n{description}", "metadata": metadata})

        if not items:
            raise DocumentProcessingError("No extractable text or visual evidence was found in this DOCX.")

        result = self._result(document_id, filename, "docx", None, items, warnings)
        result["paragraphs"] = len(doc.paragraphs)
        return result

    def _process_image(self, file_path: Path, document_id: str, filename: str) -> dict[str, Any]:
        try:
            image = Image.open(file_path)
            description = self._describe_image(image, context=f"This is standalone image file {filename}.")
        except Exception as exc:
            logger.warning("Standalone image analysis failed for document_id=%s", document_id, exc_info=exc)
            raise DocumentProcessingError("The image could not be analyzed.") from exc

        if not description:
            raise DocumentProcessingError("The image could not be converted into a visual description.")

        items = [
            {
                "text": f"Standalone image visual evidence from {filename}:\n{description}",
                "metadata": self._base_metadata(document_id, filename, SUPPORTED_EXTENSIONS[file_path.suffix.lower()], "visual"),
            }
        ]
        return self._result(document_id, filename, "image", None, items, [])

    def _extract_pdf_tables(
        self,
        file_path: Path,
        document_id: str,
        filename: str,
    ) -> tuple[list[dict[str, Any]], list[str]]:
        items: list[dict[str, Any]] = []
        warnings: list[str] = []
        try:
            import pdfplumber
        except ImportError:
            warnings.append("PDF table extraction is unavailable because pdfplumber is not installed.")
            return items, warnings

        try:
            with pdfplumber.open(str(file_path)) as pdf:
                for page_index, page in enumerate(pdf.pages, start=1):
                    tables = page.extract_tables() or []
                    for table_index, table in enumerate(tables, start=1):
                        markdown = self._table_to_markdown(table)
                        if not markdown:
                            continue
                        metadata = self._base_metadata(document_id, filename, "pdf", "table", page_index)
                        metadata["table_index"] = table_index
                        items.append(
                            {
                                "text": (
                                    f"PDF table evidence from {filename}, page {page_index}, table {table_index}. "
                                    "Use this evidence for questions about tables, rows, columns, entries, labels, "
                                    "metrics, percentages, numeric values, and captions.\n"
                                    f"{markdown}"
                                ),
                                "metadata": metadata,
                            }
                        )
        except Exception as exc:
            logger.warning("PDF table extraction failed for document_id=%s", document_id, exc_info=exc)
            warnings.append("PDF table extraction could not run for this document.")

        return items, warnings

    def _table_to_markdown(self, table: list[list[Any]]) -> str:
        rows = [
            [self._clean_cell(cell) for cell in row]
            for row in table
            if row and any(self._clean_cell(cell) for cell in row)
        ]
        if len(rows) < 2:
            return ""

        width = max(len(row) for row in rows)
        normalized = [row + [""] * (width - len(row)) for row in rows]
        header = normalized[0]
        body = normalized[1:]
        lines = [
            "| " + " | ".join(header) + " |",
            "| " + " | ".join(["---"] * width) + " |",
        ]
        lines.extend("| " + " | ".join(row) + " |" for row in body)
        return "\n".join(lines)

    def _clean_cell(self, cell: Any) -> str:
        return " ".join(str(cell or "").replace("|", "/").split())

    def _is_openai_configuration_or_quota_error(self, exc: Exception) -> bool:
        status_code = getattr(exc, "status_code", None)
        error_code = self._openai_error_code(exc)
        return status_code in {401, 402, 403, 429} or error_code in {
            "insufficient_quota",
            "billing_hard_limit_reached",
        }

    def _public_openai_visual_warning(self, exc: Exception) -> str:
        status_code = getattr(exc, "status_code", None)
        error_code = self._openai_error_code(exc)
        if status_code in {401, 403}:
            return "Visual analysis stopped because OpenAI authentication failed. Check the backend OPENAI_API_KEY."
        if status_code in {402, 429} or error_code in {"insufficient_quota", "billing_hard_limit_reached"}:
            return "Visual analysis stopped because OpenAI quota or billing is unavailable."
        return "Visual analysis stopped because OpenAI is unavailable."

    def _openai_error_code(self, exc: Exception) -> str | None:
        body = getattr(exc, "body", None)
        if isinstance(body, dict):
            error = body.get("error")
            if isinstance(error, dict) and error.get("code"):
                return str(error["code"])
            if body.get("code"):
                return str(body["code"])
        return getattr(exc, "code", None)

    def _infer_tables_from_page_text(self, text: str) -> list[str]:
        lines = [" ".join(line.split()) for line in text.splitlines() if line.strip()]
        tables: list[str] = []
        used_ranges: list[range] = []

        for caption_index, line in enumerate(lines):
            if not self._looks_like_table_caption(line):
                continue
            start = self._find_table_block_start(lines, caption_index)
            end = self._find_table_block_end(lines, caption_index)
            block_range = range(start, end)
            if any(set(block_range).intersection(existing) for existing in used_ranges):
                continue
            block = lines[start:end]
            if self._looks_like_table_block(block):
                tables.append("\n".join(block))
                used_ranges.append(block_range)

        return tables

    def _looks_like_table_caption(self, line: str) -> bool:
        return bool(re.match(r"^\s*(table|tab\.)\s*\d+\s*[:.\-]", line, flags=re.IGNORECASE))

    def _find_table_block_start(self, lines: list[str], caption_index: int) -> int:
        start = caption_index
        saw_data = False
        for index in range(caption_index - 1, max(-1, caption_index - 12), -1):
            if self._looks_like_table_data_line(lines[index]):
                start = index
                saw_data = True
                continue
            if saw_data and self._looks_like_header_line(lines[index]):
                start = index
                break
            if saw_data:
                break
        return start

    def _find_table_block_end(self, lines: list[str], caption_index: int) -> int:
        end = caption_index + 1
        saw_data = False
        for index in range(caption_index + 1, min(len(lines), caption_index + 12)):
            if self._looks_like_table_data_line(lines[index]):
                end = index + 1
                saw_data = True
                continue
            if not saw_data and self._looks_like_header_line(lines[index]):
                end = index + 1
                continue
            if saw_data:
                break
        return end

    def _looks_like_table_block(self, lines: list[str]) -> bool:
        if len(lines) < 3:
            return False
        data_like = sum(1 for line in lines if self._looks_like_table_data_line(line))
        has_caption = any(self._looks_like_table_caption(line) for line in lines)
        has_header = any(self._looks_like_header_line(line) for line in lines[:3])
        return data_like >= 2 and (has_caption or has_header)

    def _looks_like_table_data_line(self, line: str) -> bool:
        if len(line) > 180:
            return False
        numeric_tokens = self._count_numeric_tokens(line)
        words = re.findall(r"[A-Za-z][A-Za-z0-9_()/-]*", line)
        if numeric_tokens >= 2:
            return True
        if "%" in line and numeric_tokens >= 1 and len(words) >= 1:
            return True
        return False

    def _looks_like_header_line(self, line: str) -> bool:
        words = re.findall(r"[A-Za-z][A-Za-z0-9_()/-]*", line)
        if len(words) < 2 or len(line) > 140:
            return False
        if line.endswith("."):
            return False
        numeric_tokens = self._count_numeric_tokens(line)
        return numeric_tokens == 0 and len(words) <= 10

    def _count_numeric_tokens(self, line: str) -> int:
        return len(re.findall(r"[-+]?\d+(?:[.,]\d+)?%?", line))

    def _infer_table_from_page_text(self, text: str) -> str:
        tables = self._infer_tables_from_page_text(text)
        return tables[0] if tables else ""

    def _describe_image(self, image: Image.Image, context: str = "") -> str:
        client = self._client()
        data_url = self._image_to_data_url(image)
        prompt = (
            f"{context}\n\n"
            "Return a concise but literal Markdown description for retrieval. Do not summarize away details. "
            "Extract exact visible text, labels, captions, figure numbers, table numbers, legends, axes, row/column "
            "headers, ranks, team names, scores, percentages, and numeric values. If a pie chart or chart is visible, "
            "list every visible category and percentage/value exactly. If a table is visible, transcribe it as a "
            "Markdown table when readable. If a value is not readable, write 'unreadable' instead of guessing. "
            "Prefer this structure:\n"
            "- Visible identifiers:\n"
            "- Figures/charts:\n"
            "- Tables:\n"
            "- Percentages and numeric values:\n"
            "- Important text:"
        )
        response = client.responses.create(
            model=OPENAI_MODEL,
            input=[
                {
                    "role": "user",
                    "content": [
                        {"type": "input_text", "text": prompt},
                        {
                            "type": "input_image",
                            "image_url": data_url,
                            "detail": OPENAI_IMAGE_DETAIL,
                        },
                    ],
                }
            ],
            max_output_tokens=VISUAL_DESCRIPTION_MAX_TOKENS,
        )
        return (getattr(response, "output_text", "") or "").strip()

    def _client(self) -> OpenAI:
        if self.client is None:
            self.client = OpenAI(api_key=require_openai_api_key())
        return self.client

    def _image_to_data_url(self, image: Image.Image) -> str:
        with tempfile.SpooledTemporaryFile() as buffer:
            clean_image = image.convert("RGB")
            clean_image.thumbnail((2400, 2400))
            clean_image.save(buffer, format="JPEG", quality=85)
            buffer.seek(0)
            encoded = base64.b64encode(buffer.read()).decode("ascii")
        return f"data:image/jpeg;base64,{encoded}"

    def _result(
        self,
        document_id: str,
        filename: str,
        file_type: str,
        pages: int | None,
        items: list[dict[str, Any]],
        warnings: list[str],
    ) -> dict[str, Any]:
        return {
            "document": {
                "id": document_id,
                "filename": filename,
                "file_type": file_type,
                "pages": pages,
                "text_items": sum(1 for item in items if item["metadata"]["content_type"] == "text"),
                "visual_items": sum(1 for item in items if item["metadata"]["content_type"] in {"visual", "table"}),
            },
            "items": items,
            "warnings": warnings,
        }
